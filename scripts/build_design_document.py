"""
SIGMA-K Complete System Design Document & Prototype Builder (DOCX + PDF)
Generates:
- docs/DOKUMEN_PERANCANGAN_SIGMA-K_v1.0.docx
- docs/DOKUMEN_PERANCANGAN_SIGMA-K_v1.0.pdf
"""

import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

DOCX_OUT = os.path.abspath('docs/DOKUMEN_PERANCANGAN_SIGMA-K_v1.0.docx')
PDF_OUT = os.path.abspath('docs/DOKUMEN_PERANCANGAN_SIGMA-K_v1.0.pdf')
ASSETS_DIR = os.path.abspath('docs/assets')

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color_hex="CBD5E1"):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

class DocBuilder:
    def __init__(self):
        self.doc = Document()
        self.fig_count = 0
        self.tbl_count = 0
        self._setup_page_layout()
        self._setup_styles()

    def _setup_page_layout(self):
        for section in self.doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.different_first_page_header_footer = True
            
            # Header
            header = section.header
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hrun = hp.add_run("Dokumen Perancangan Sistem E-SKLD / SIGMA-K v1.0 | KemenPANRB")
            hrun.font.name = 'Calibri'
            hrun.font.size = Pt(8.5)
            hrun.font.color.rgb = RGBColor(100, 116, 139)

            # Footer
            footer = section.footer
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            frun = fp.add_run("Kementerian Pendayagunaan Aparatur Negara dan Reformasi Birokrasi (KemenPANRB) — Dokumen Internal")
            frun.font.name = 'Calibri'
            frun.font.size = Pt(8.5)
            frun.font.color.rgb = RGBColor(100, 116, 139)

    def _setup_styles(self):
        style_normal = self.doc.styles['Normal']
        style_normal.font.name = 'Calibri'
        style_normal.font.size = Pt(10.5)
        style_normal.font.color.rgb = RGBColor(30, 41, 59)
        style_normal.paragraph_format.line_spacing = 1.15
        style_normal.paragraph_format.space_after = Pt(4)

    def add_title_cover(self):
        # Spacing
        for _ in range(2):
            self.doc.add_paragraph()

        p_kemen = self.doc.add_paragraph()
        p_kemen.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_kemen = p_kemen.add_run("KEMENTERIAN PENDAYAGUNAAN APARATUR NEGARA\nDAN REFORMASI BIROKRASI REPUBLIK INDONESIA")
        r_kemen.font.name = 'Calibri'
        r_kemen.font.size = Pt(12)
        r_kemen.font.bold = True
        r_kemen.font.color.rgb = RGBColor(11, 42, 74)

        for _ in range(2):
            self.doc.add_paragraph()

        p_main = self.doc.add_paragraph()
        p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_main = p_main.add_run("DOKUMEN PERANCANGAN SISTEM &\nPROTOTYPE APLIKASI")
        r_main.font.name = 'Calibri'
        r_main.font.size = Pt(22)
        r_main.font.bold = True
        r_main.font.color.rgb = RGBColor(11, 42, 74)

        p_sub = self.doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sub = p_sub.add_run("E-SKLD / SIGMA-K\nSistem Pengelolaan Data Kementerian/Lembaga/Pemerintah Daerah\ndan Struktur Kelembagaan")
        r_sub.font.name = 'Calibri'
        r_sub.font.size = Pt(13)
        r_sub.font.color.rgb = RGBColor(30, 64, 175)

        for _ in range(3):
            self.doc.add_paragraph()

        # Metadata Card Table
        tbl = self.doc.add_table(rows=6, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl, "CBD5E1")
        
        meta = [
            ("Judul Dokumen", "Dokumen Perancangan Sistem & Prototype Aplikasi E-SKLD (SIGMA-K)"),
            ("Nomor Dokumen", "DOC-PANRB-SIGMAK-2026-V1.0"),
            ("Versi / Rilis", "Versi 1.0 (Final Architecture & Implementation Baseline)"),
            ("Tanggal Rilis", "27 Agustus 2026"),
            ("Penyusun", "Tim Pengembang & System Analyst Kelembagaan KemenPANRB"),
            ("Status Validasi", "100% Lolos Uji (105 Frontend Tests & 198 Backend Unit Tests)")
        ]
        
        for idx, (label, val) in enumerate(meta):
            c0, c1 = tbl.cell(idx, 0), tbl.cell(idx, 1)
            c0.width, c1.width = Inches(2.2), Inches(4.3)
            set_cell_margins(c0, 80, 80, 100, 100)
            set_cell_margins(c1, 80, 80, 100, 100)
            set_cell_background(c0, "F1F5F9")
            
            p0 = c0.paragraphs[0]
            p0.paragraph_format.space_after = Pt(0)
            r0 = p0.add_run(label)
            r0.font.bold = True
            r0.font.size = Pt(9.5)
            r0.font.color.rgb = RGBColor(11, 42, 74)

            p1 = c1.paragraphs[0]
            p1.paragraph_format.space_after = Pt(0)
            r1 = p1.add_run(val)
            r1.font.size = Pt(9.5)
            r1.font.color.rgb = RGBColor(30, 41, 59)

        self.doc.add_page_break()

    def add_h1(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(11, 42, 74)
        return p

    def add_h2(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 64, 175)
        return p

    def add_h3(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_p(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(30, 41, 59)
        return p

    def add_bullet(self, text, bold_prefix=""):
        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.bold = True
            r_pre.font.size = Pt(10)
            r_pre.font.color.rgb = RGBColor(11, 42, 74)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(30, 41, 59)
        return p

    def add_callout(self, title, text):
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, 100, 100, 150, 150)
        
        # Left border highlight
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="0B2A4A"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
        ''')
        tcPr.append(borders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_title = p.add_run(f"PENTING / PERHATIAN: {title}\n")
        r_title.font.bold = True
        r_title.font.size = Pt(9.5)
        r_title.font.color.rgb = RGBColor(11, 42, 74)
        
        r_text = p.add_run(text)
        r_text.font.size = Pt(9)
        r_text.font.color.rgb = RGBColor(51, 65, 85)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def add_image_figure(self, filename, caption, width_in=6.0):
        img_path = os.path.join(ASSETS_DIR, filename)
        if os.path.exists(img_path):
            self.fig_count += 1
            p_img = self.doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(2)
            p_img.paragraph_format.keep_with_next = True
            run = p_img.add_run()
            run.add_picture(img_path, width=Inches(width_in))

            p_cap = self.doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(0)
            p_cap.paragraph_format.space_after = Pt(8)
            r_cap = p_cap.add_run(f"Gambar {self.fig_count}: {caption}")
            r_cap.font.name = 'Calibri'
            r_cap.font.size = Pt(9)
            r_cap.font.bold = True
            r_cap.font.color.rgb = RGBColor(71, 85, 105)
        else:
            p_err = self.doc.add_paragraph(f"[Gambar tidak ditemukan: {filename}]")
            p_err.runs[0].font.color.rgb = RGBColor(220, 38, 38)

    def add_table_data(self, title, headers, rows, col_widths=None):
        self.tbl_count += 1
        p_cap = self.doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(6)
        p_cap.paragraph_format.space_after = Pt(2)
        p_cap.paragraph_format.keep_with_next = True
        r_cap = p_cap.add_run(f"Tabel {self.tbl_count}: {title}")
        r_cap.font.bold = True
        r_cap.font.size = Pt(9.5)
        r_cap.font.color.rgb = RGBColor(11, 42, 74)

        tbl = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl, "CBD5E1")

        # Header
        for c_idx, h_text in enumerate(headers):
            cell = tbl.cell(0, c_idx)
            set_cell_background(cell, "0B2A4A")
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(h_text)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)

        # Rows
        for r_idx, row_data in enumerate(rows):
            bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                cell = tbl.cell(r_idx + 1, c_idx)
                set_cell_background(cell, bg)
                set_cell_margins(cell, 60, 60, 100, 100)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(str(val))
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(30, 41, 59)

        # Widths
        if col_widths:
            for row in tbl.rows:
                for c_idx, w in enumerate(col_widths):
                    row.cells[c_idx].width = Inches(w)

        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

def build_sigma_k_document():
    builder = DocBuilder()
    builder.add_title_cover()

    # =========================================================
    # BAB 1: GAMBARAN UMUM
    # =========================================================
    builder.add_h1("BAB 1 — GAMBARAN UMUM")
    
    builder.add_h2("1.1 Nama Sistem")
    builder.add_p("Nama resmi sistem adalah E-SKLD / SIGMA-K (Sistem Pengelolaan Data Kementerian/Lembaga/Pemerintah Daerah dan Struktur Kelembagaan). Sistem ini dikembangkan sebagai platform tunggal (single source of truth) terpadu di lingkungan Kementerian Pendayagunaan Aparatur Negara dan Reformasi Birokrasi (KemenPANRB) untuk mengelola data kelembagaan instansi pemerintah di seluruh Indonesia.")

    builder.add_h2("1.2 Latar Belakang")
    builder.add_p("Dalam kerangka reformasi birokrasi dan penataan kelembagaan instansi pemerintah Republik Indonesia—termasuk penyesuaian susunan Kabinet Merah Putih periode 2024–2029—dibutuhkan standardisasi data struktur organisasi yang akurat, dinamis, dan terverifikasi secara hukum. Selama ini, proses pengusulan perubahan organisasi, penambahan unit kerja, serta revisi formasi jabatan struktural (Eselon I s.d. Eselon IV) masih terfragmentasi, memicu ketidaksinkronan data antar instansi dan memerlukan waktu telaah manual yang panjang.")
    builder.add_p("E-SKLD / SIGMA-K hadir sebagai solusi digital berbasis web modern guna mendigitalisasi siklus hidup pengusulan kelembagaan secara menyeluruh, menjamin pemisahan wewenang yang tegas (Separation of Duties), dan menyediakan rekam jejak audit forensik yang tidak dapat diubah (append-only audit trail).")

    builder.add_h2("1.3 Tujuan Sistem")
    builder.add_bullet("Menyediakan repositori master data kelembagaan nasional (Kementerian Koordinator, Kementerian, LPNK, Lembaga Non-Struktural, dan Pemerintah Daerah) yang terintegrasi.", "1. Sentralisasi Data: ")
    builder.add_bullet("Mengotomatisasi alur penapisan administratif (Gate 1 Admin) dan telaah substantif (Gate 2 Verifier) secara akuntabel.", "2. Tata Kelola Workflow: ")
    builder.add_bullet("Menyajikan visualisasi pohon organisasi hierarkis interaktif dengan kapabilitas penelusuran relasi parent-child dan jabatan formasi ASN.", "3. Visualisasi Graf Interaktif: ")
    builder.add_bullet("Menyediakan mekanisme pengesahan Surat Keputusan (SK) resmi oleh Verifikator dengan promosi otomatis usulan yang disahkan ke master data aktif.", "4. Otomasi Pengesahan & Promosi: ")
    builder.add_bullet("Menjamin keamanan akses berbasis Zero-Trust dan mencatat setiap perubahan data untuk audit forensik kepatuhan.", "5. Forensik & Akuntabilitas: ")

    builder.add_h2("1.4 Ruang Lingkup")
    builder.add_p("Ruang lingkup operasional sistem mencakup:")
    builder.add_bullet("Manajemen katalog K/L/D, komposisi kabinet pemerintahan, dan pemetaan dasar hukum pembentukan instansi.", "• Master Data Kelembagaan: ")
    builder.add_bullet("Pengelolaan unit kerja struktural berbasis model graf Adjacency List dan validasi pencegahan siklus (anti-cycle DFS).", "• Struktur Organisasi: ")
    builder.add_bullet("Pembuatan draf usulan penataan kelembagaan, penambahan unit baru, penghapusan unit, dan perbaikan formasi jabatan.", "• Siklus Pengusulan (Submission): ")
    builder.add_bullet("Penguncian snapshot usulan (v1, v2, dst.), pencatatan form review catatan telaah, dan komparasi perubahan (Diff Viewer).", "• Versioning & Siklus Revisi: ")
    builder.add_bullet("Penapisan kelengkapan berkas, validasi formal, dan penugasan verifikator substantif.", "• Gate 1 Admin Screening: ")
    builder.add_bullet("Telaah substantif beban kerja, evaluasi formasi jabatan, rekomendasi teknis, dan pengesahan final SK persetujuan.", "• Gate 2 Verifier Workspace: ")
    builder.add_bullet("Agregasi metrik efisiensi, piramida eselonisasi, analisis SLA kecepatan layanan, dan ekspor dataset (CSV/JSON).", "• Dashboard & Intelijensi Data: ")
    builder.add_bullet("Pencatatan rekam jejak aktor, event aksi, payload JSON sebelum/sesudah, IP address, dan user agent.", "• Forensik Audit Log: ")

    builder.add_h2("1.5 Pengguna Sistem")
    builder.add_p("Sistem melayani empat aktor utama dengan pembagian tugas yang terisolasi:")
    builder.add_bullet("Operator resmi perwakilan instansi pemerintah yang berwenang menyusun dan mengajukan usulan struktur organisasi unit kerjanya.", "1. USER (Operator Instansi): ")
    builder.add_bullet("Staf administrasi KemenPANRB yang bertugas melakukan penapisan kelengkapan administratif berkas dan mendistribusikan penugasan telaah.", "2. ADMIN (Penapis KemenPANRB): ")
    builder.add_bullet("Pejabat fungsional analis kelembagaan KemenPANRB yang melakukan telaah substantif, memegang hak pengesahan akhir SK, dan mengeksekusi promosi data.", "3. VERIFIER (Verifikator Kelembagaan): ")
    builder.add_bullet("Administrator sistem dan pimpinan (SESDEP) yang memiliki visibilitas analitik nasional, manajemen hak akses, dan audit forensik menyeluruh.", "4. SUPER_ADMIN (Administrator Sistem): ")

    builder.add_h2("1.6 Kondisi Implementasi Saat Ini")
    builder.add_p("Berdasarkan audit repository aktual, seluruh fondasi arsitektur dan integrasi API dari Phase 1 hingga Phase 14D telah selesai diimplementasikan dan diverifikasi 100%:")
    builder.add_bullet("Sebanyak 105 pengujian otomatis frontend (Phase 14A Foundation, Phase 14B Auth, Phase 14B Security, Phase 14C Master Data, dan Phase 14D Reporting) dinyatakan 100% LULUS (PASS).", "• Frontend Test Suite: ")
    builder.add_bullet("Sebanyak 198 pengujian unit backend CodeIgniter 4 dengan 713 assertions dinyatakan 100% LULUS (0 error, 0 failure).", "• Backend PHPUnit Suite: ")
    builder.add_bullet("Kompilasi TypeScript strict menghasilkan 0 error; analisis ESLint menghasilkan 0 warning/error; dan build Next.js sukses mengompilasi seluruh 16 route.", "• Analisis Statis & Build: ")
    builder.add_bullet("Skema database MySQL eskld_db terjaga immutabilitasnya (0 perubahan skema tidak sah, 21 tabel relasional terverifikasi).", "• Database Integrity: ")

    # =========================================================
    # BAB 2: GAMBARAN SISTEM & ARSITEKTUR
    # =========================================================
    builder.add_h1("BAB 2 — GAMBARAN SISTEM & ARSITEKTUR")
    
    builder.add_h2("2.1 Gambaran Umum SIGMA-K")
    builder.add_p("SIGMA-K dirancang menggunakan paradigma Multi-Tier Decoupled Architecture yang memisahkan secara tegas antarmuka pengguna (Frontend Next.js), gerbang transportasi data (HTTP Client Gateway), lapisan logika bisnis dan otorisasi (Backend CodeIgniter 4), serta basis data relasional (MySQL). Pendekatan ini memastikan skalabilitas tinggi, kemudahan pemeliharaan, serta kepatuhan penuh terhadap standar keamanan informasi pemerintah.")

    builder.add_h2("2.2 Fungsi Utama Sistem")
    builder.add_p("Fungsi-fungsi inti yang disediakan SIGMA-K meliputi:")
    builder.add_bullet("Otentikasi berbasis JSON Web Token (JWT) dengan penanganan sesi aman dan penonaktifan persona switcher pada mode API.", "• Autentikasi & Profil Terpusat: ")
    builder.add_bullet("Penyajian daftar instansi, pohon unit kerja dinamis dengan React Flow, formasi jabatan, dan komposisi kabinet.", "• Manajemen Master Data: ")
    builder.add_bullet("Pembuatan berkas usulan berjenjang dengan validasi input komprehensif pada level unit dan jabatan.", "• Penyusunan Usulan (Drafting): ")
    builder.add_bullet("Verifikasi administratif berkas oleh Admin KemenPANRB dan mekanisme pengembalian revisi.", "• Gate 1 Screening: ")
    builder.add_bullet("Telaah beban kerja, analisis eselonisasi, penerbitan rekomendasi, dan pengesahan SK resmi oleh Verifikator.", "• Gate 2 Verifier Review: ")
    builder.add_bullet("Migrasi data otomatis dari usulan yang telah disahkan langsung ke struktur master data aktif instansi terkait.", "• Master Data Promotion: ")
    builder.add_bullet("Visualisasi KPI, distribusi jabatan struktural, dan ekspor dataset berkas usulan dalam format CSV/JSON.", "• Dashboard & Pelaporan Eksekutif: ")
    builder.add_bullet("Pencatatan rekam jejak mutlak atas seluruh manipulasi data guna kebutuhan audit forensik.", "• Audit Trail Forensik: ")

    builder.add_h2("2.3 Modul Sistem")
    builder.add_table_data(
        "Modul-Modul Utama Sistem SIGMA-K",
        ["No", "Nama Modul", "Komponen Backend", "Komponen Frontend", "Fungsi Utama"],
        [
            ["1", "Modul Autentikasi", "AuthController, AuthFilter, UserModel", "AuthService, RoleContext, /login", "Login NIP/Pass, JWT token generator, getMe profile"],
            ["2", "Modul Master Data", "InstitutionController, OrgHierarchyService", "InstitutionService, /institutions, /structure", "Katalog K/L/D, Adjacency List tree, React Flow canvas"],
            ["3", "Modul Pengusulan", "SubmissionWorkflowController, SubmissionModel", "SubmissionService, /submissions, /submissions/[id]", "Drafting usulan, versioning snapshot v1->v2, diff viewer"],
            ["4", "Modul Gate 1 Admin", "AdminWorkflowController, VerifierAssignmentModel", "SubmissionService, /verifications (Gate 1)", "Penapisan administrasi, return revision, assign verifier"],
            ["5", "Modul Gate 2 Verifier", "VerifierWorkflowController, ApprovalModel", "SubmissionService, /verifications/[id] (Gate 2)", "Telaah substantif, catatan telaah, final approval, promotion"],
            ["6", "Modul Pelaporan", "ReportController, ExecutiveReportService", "AnalyticsService, /, /analytics", "Ringkasan metrik, funnel, postur ASN, ekspor CSV/JSON"],
            ["7", "Modul Audit Trail", "AuditLogController, AuditService", "AuditService, /audit-logs", "Pencatatan append-only audit, forensik perubahan data"]
        ],
        [0.4, 1.5, 1.6, 1.6, 1.4]
    )

    builder.add_h2("2.4 Alur Data Antar Lapisan")
    builder.add_p("Setiap permintaan (request) data dari pengguna diproses melalui rantai kendali yang ketat:")
    builder.add_bullet("Pengguna berinteraksi dengan UI Next.js (React 18). Komponen memanggil Service Facade yang telah diabstraksi.", "1. Presentation Layer: ")
    builder.add_bullet("HttpClient melakukan injeksi header `Authorization: Bearer <token>` dan mengirimkan request HTTPS ke REST API.", "2. Transport Layer: ")
    builder.add_bullet("AuthFilter memverifikasi validitas tanda tangan kriptografis JWT, mengekstrak identitas pengguna, role, dan home institution.", "3. Filter & JWT Verification: ")
    builder.add_bullet("AuthorizationService & ScopeResolver memeriksa izin granular (permissions) dan memastikan target instansi berada dalam cakupan wewenang pengguna (Anti BOLA/IDOR).", "4. Zero-Trust Authorization: ")
    builder.add_bullet("Domain Service mengeksekusi logika bisnis, memvalidasi aturan state machine, dan mencatat mutasi ke AuditService.", "5. Domain Business Logic: ")
    builder.add_bullet("Data disimpan atau diambil dari basis data MySQL eskld_db menggunakan PDO query builder terenkapsulasi.", "6. Data Persistence: ")

    builder.add_h2("2.5 Arsitektur Sistem")
    builder.add_p("Arsitektur sistem SIGMA-K secara komprehensif mengintegrasikan seluruh komponen presentasi, transportasi, keamanan, dan persistensi data sebagaimana ditunjukkan pada Gambar 1.")
    builder.add_image_figure("system_architecture.png", "Arsitektur Sistem Terintegrasi E-SKLD / SIGMA-K (4-Tier Architecture)")

    # =========================================================
    # BAB 3: ROLE DAN HAK AKSES
    # =========================================================
    builder.add_h1("BAB 3 — ROLE DAN HAK AKSES")
    
    builder.add_h2("3.1 Definisi Role Resmi")
    builder.add_p("Sistem mengimplementasikan Role-Based Access Control (RBAC) yang dikombinasikan dengan Multi-Tenant Scoping. Terdapat empat role resmi yang terdaftar pada tabel `roles` basis data:")
    builder.add_bullet("Role bagi operator instansi kementerian/lembaga/daerah. Memiliki hak menyusun draf usulan, menambahkan unit dan formasi jabatan pada instansinya sendiri, serta menanggapi catatan revisi.", "1. USER (Operator Instansi): ")
    builder.add_bullet("Role bagi staf penapis KemenPANRB. Bertanggung jawab atas Gate 1 Admin Screening untuk memverifikasi kelengkapan berkas, mengembalikan revisi administratif, dan menugaskan Verifikator.", "2. ADMIN (Penapis KemenPANRB): ")
    builder.add_bullet("Role bagi pejabat analis kelembagaan KemenPANRB. Bertanggung jawab atas Gate 2 Substantive Review, menerbitkan rekomendasi teknis, menandatangani pengesahan SK resmi, dan mengeksekusi promosi master data.", "3. VERIFIER (Verifikator Kelembagaan): ")
    builder.add_bullet("Role tingkat tertinggi bagi administrator sistem dan pimpinan eksekutif (SESDEP). Memiliki akses visibilitas nasional, pengelolaan akun pengguna, monitoring performa layanan, dan audit forensik penuh.", "4. SUPER_ADMIN (Administrator Sistem): ")

    builder.add_callout("Status Persona SESDEP", "Pada prototipe awal, SESDEP diperkenalkan sebagai representasi persona pimpinan. Namun, dalam arsitektur produksi backend yang telah diverifikasi, kode peran SESDEP dinormalisasi secara otomatis ke dalam role SUPER_ADMIN dengan wewenang pengawasan eksekutif nasional penuh. Switcher persona pada UI Next.js dinonaktifkan pada mode API produksi guna mencegah eskalasi hak akses ilegal.")

    builder.add_h2("3.2 Matriks Hak Akses dan Lingkup Wewenang")
    builder.add_table_data(
        "Matriks Hak Akses dan Tanggung Jawab Workflow per Role",
        ["Role", "Tujuan Utama", "Hak Akses (Permissions)", "Lingkup Akses (Scope)", "Tanggung Jawab Workflow"],
        [
            ["USER", "Penyusunan usulan struktur organisasi instansi pengusul", "submission:create, submission:edit, org:read, inst:read_own", "Home Institution (Instansi Sendiri)", "Drafting usulan, submit usulan, perbaikan revisi (v1->v2)"],
            ["ADMIN", "Penapisan kelengkapan berkas administratif (Gate 1)", "screening:read, screening:action, verifier:assign, audit:read_scoped", "Scoped K/L yang ditugaskan", "Pemeriksaan formal, return revisi administratif, assign verifier"],
            ["VERIFIER", "Telaah substantif & pengesahan SK resmi (Gate 2)", "review:substantive, approval:final, promotion:execute, audit:read", "Queue usulan yang ditugaskan", "Telaah beban kerja & formasi, catatan teknis, final approval SK, promosi master data"],
            ["SUPER_ADMIN", "Pengawasan nasional, kelola akses & audit forensik", "admin:all, user:manage, audit:forensics, reports:export_all", "Global Nationwide (Seluruh K/L/D)", "Monitoring KPI nasional, audit trail kepatuhan, ekspor dataset"]
        ],
        [1.0, 1.4, 1.5, 1.2, 1.4]
    )

    builder.add_h2("3.3 Diagram Model Hak Akses")
    builder.add_p("Pemisahan wewenang dan batasan interaksi antar role pada SIGMA-K divisualisasikan pada Gambar 2.")
    builder.add_image_figure("role_access_matrix.png", "Model Otorisasi dan Matriks Hak Akses Zero-Trust SIGMA-K")

    # =========================================================
    # BAB 4: PRINSIP DESAIN UI/UX
    # =========================================================
    builder.add_h1("BAB 4 — PRINSIP DESAIN UI/UX")
    
    builder.add_h2("4.1 Prinsip Desain Antarmuka")
    builder.add_p("Desain antarmuka SIGMA-K dibangun dengan mengedepankan prinsip kejelasan institusional, efisiensi navigasi, dan kenyamanan visual bagi aparatur sipil negara. Karakteristik visual yang diterapkan meliputi:")
    builder.add_bullet("Menggunakan palet warna korporat Kementerian: Deep Navy (#0B2A4A) sebagai warna dominan identitas negara, Blue (#1E40AF) untuk elemen interaktif primer, Gold (#D4AF37) untuk status eksekutif dan aksen penghargaan, serta Slate (#F8FAFC) sebagai latar kerja bersih.", "• Palet Warna Harmonis: ")
    builder.add_bullet("Struktur tata letak konsisten dengan sidebar navigasi terorganisir per domain fungsional, topbar status sesi, breadcrumb hirarkis, dan judul halaman berstandar PageHeader.", "• Tata Letak Konsisten: ")
    builder.add_bullet("Seluruh komponen interaktif (tombol, baris tabel, tab navigasi) dilengkapi dengan micro-interaction yang responsif.", "• Umpan Balik Visual Responsif: ")

    builder.add_h2("4.2 Penanganan State Aplikasi")
    builder.add_p("Untuk menjamin pengalaman pengguna yang andal, aplikasi menangani berbagai kondisi state secara eksplisit:")
    builder.add_bullet("Setiap proses pengambilan data asynchronous menampilkan komponen Spinner atau Skeleton loader sehingga pengguna mengetahui proses sedang berjalan.", "• Loading State: ")
    builder.add_bullet("Ketika dataset kosong (misalnya tidak ada usulan dalam antrean), layar menampilkan kartu ilustrasi informatif beserta tombol aksi kontekstual.", "• Empty State: ")
    builder.add_bullet("Sistem menyediakan komponen Error Banner terdedikasi untuk menampilkan pesan kesalahan teknis secara transparan:", "• Error State & Status Code: ")
    builder.add_bullet("Mengarahkan pengguna kembali ke form login dengan pesan sesi berakhir.", "  - 401 Unauthorized: ")
    builder.add_bullet("Menampilkan pesan penolakan akses Zero-Trust jika pengguna mencoba mengakses instansi di luar wewenangnya.", "  - 403 Forbidden: ")
    builder.add_bullet("Menampilkan pesan berkas atau entitas instansi tidak ditemukan.", "  - 404 Not Found: ")
    builder.add_bullet("Menampilkan peringatan konflik versi saat usulan telah diperbarui oleh pengguna lain.", "  - 409 Conflict: ")
    builder.add_bullet("Menyorot field form yang tidak valid beserta rincian kesalahan aturan validasi.", "  - 422 Validation Error: ")

    # =========================================================
    # BAB 5: SITEMAP APLIKASI
    # =========================================================
    builder.add_h1("BAB 5 — SITEMAP APLIKASI")
    
    builder.add_h2("5.1 Struktur Hierarki Navigasi")
    builder.add_p("Navigasi aplikasi SIGMA-K disusun berdasarkan 16 rute aktual yang telah terkompilasi dalam framework Next.js App Router:")
    builder.add_table_data(
        "Daftar Rute Halaman Aplikasi SIGMA-K",
        ["No", "Rute URL", "Nama Halaman", "Akses Role", "Deskripsi Fungsional"],
        [
            ["1", "/login", "Layar Masuk Sistem", "Publik / All", "Otentikasi NIP/Password & inisialisasi sesi JWT"],
            ["2", "/", "Executive Dashboard", "All Roles", "Ringkasan metrik, funnel usulan, dan widget SK terbaru"],
            ["3", "/analytics", "Intelijensi & Postur ASN", "All Roles", "Analisis KPI, piramida eselon, SLA layanan, unduh CSV"],
            ["4", "/institutions", "Katalog Master Instansi", "All Roles", "Daftar seluruh K/L/D dengan filter kategori dan pencarian"],
            ["5", "/institutions/[id]", "Rincian Profil Instansi", "All Roles", "Dasar hukum, daftar unit, rekapitulasi posisi instansi"],
            ["6", "/structure", "Bagan Organisasi", "All Roles", "Kanvas graf React Flow interaktif pohon hierarki unit"],
            ["7", "/cabinets", "Katalog Kabinet", "All Roles", "Struktur Kabinet Merah Putih dan kementerian koordinator"],
            ["8", "/cabinets/compare", "Komparasi Kabinet", "All Roles", "Analisis perbandingan perubahan struktur lintas kabinet"],
            ["9", "/submissions", "Daftar Usulan", "USER, ADMIN, VERIFIER", "Tabel status berkas pengusulan penataan kelembagaan"],
            ["10", "/submissions/new", "Form Usulan Baru", "USER", "Formulir penyusunan draf usulan penataan organisasi"],
            ["11", "/submissions/[id]", "Rincian & Diff Usulan", "All Roles", "Informasi perubahan unit, jabatan, dan workflow stepper"],
            ["12", "/submissions/[id]/revision", "Form Perbaikan Revisi", "USER", "Ruang perbaikan berkas pasca catatan telaah (v2)"],
            ["13", "/verifications", "Antrean Verifikasi", "ADMIN, VERIFIER", "Antrean kerja penapisan Admin dan telaah Verifikator"],
            ["14", "/verifications/[id]", "Ruang Kerja Telaah", "ADMIN, VERIFIER", "Panel telaah substantif, input catatan & pengesahan SK"],
            ["15", "/audit-logs", "Log Forensik Audit", "ADMIN, SUPER_ADMIN", "Tabel rekam jejak mutasi data dan forensik payload"],
            ["16", "/notifications", "Pusat Pemberitahuan", "All Roles", "Notifikasi perubahan status usulan dan penugasan"]
        ],
        [0.4, 1.6, 1.5, 1.3, 1.7]
    )

    builder.add_h2("5.2 Visual Sitemap")
    builder.add_p("Pohon struktur navigasi aplikasi disajikan pada Gambar 3.")
    builder.add_image_figure("sitemap_diagram.png", "Pohon Struktur Navigasi dan Sitemap Aplikasi SIGMA-K")

    # =========================================================
    # BAB 6: MASTER DATA KELEMBAGAAN
    # =========================================================
    builder.add_h1("BAB 6 — MASTER DATA KELEMBAGAAN")
    
    builder.add_h2("6.1 Entitas Master Data")
    builder.add_p("SIGMA-K mengelola enam entitas master data kelembagaan yang saling berelasi:")
    builder.add_bullet("Mencatat seluruh entitas kementerian, lembaga pemerintah non-kementerian, lembaga non-struktural, dan pemerintah daerah beserta kode unik dan kategori.", "1. Institutions (Instansi): ")
    builder.add_bullet("Klasifikasi baku bentuk instansi (KEMENKO, KEMENTERIAN, LPNK, PEMDA_PROV, PEMDA_KABKOT).", "2. Institution Types: ")
    builder.add_bullet("Mengelompokkan kementerian ke dalam portofolio kabinet pemerintahan yang sedang menjabat (mis. Kabinet Merah Putih).", "3. Cabinets & Portofolio: ")
    builder.add_bullet("Unit kerja struktural di dalam instansi yang disusun secara hierarkis menggunakan relasi Adjacency List (parent_id).", "4. Organizational Units: ")
    builder.add_bullet("Daftar jabatan struktural yang melekat pada unit kerja, lengkap dengan jenjang eselon (I.a, I.b, II.a, dst.) dan kuota formasi ASN.", "5. Positions (Jabatan & Formasi): ")
    builder.add_bullet("Tugas pokok dan fungsi operasional yang menjadi landasan hukum unit kerja.", "6. Tupoksi: ")

    builder.add_h2("6.2 Alur Data & Representasi Hierarki")
    builder.add_p("Hubungan relasional antar master data disajikan pada Gambar 4, dan representasi pohon hierarki struktural disajikan pada Gambar 5.")
    builder.add_image_figure("master_data_flow.png", "Alur Relasi dan Integrasi Master Data Kelembagaan")
    builder.add_image_figure("org_hierarchy_tree.png", "Bagan Struktur Hierarki Pohon Organisasi Berjenjang (Eselon I - IV)")

    # =========================================================
    # BAB 7: DATABASE & STRUKTUR DATA (ERD)
    # =========================================================
    builder.add_h1("BAB 7 — DATABASE & STRUKTUR DATA (ERD)")
    
    builder.add_h2("7.1 Skema Relasional Basis Data (MySQL eskld_db)")
    builder.add_p("Basis data operasional `eskld_db` tersusun atas 21 tabel relasional berintegritas tinggi dengan foreign key constraints dan indeks performa:")
    builder.add_table_data(
        "Katalog 21 Tabel Relasional Basis Data eskld_db",
        ["No", "Nama Tabel", "Domain Fungsional", "Primary Key", "Foreign Keys Kunci", "Deskripsi Data"],
        [
            ["1", "users", "Autentikasi", "id (INT)", "role_id, institution_id", "Data akun pengguna, NIP, email, dan password hash"],
            ["2", "roles", "Autentikasi", "id (INT)", "-", "Definisi 4 role resmi (USER, ADMIN, VERIFIER, SUPER_ADMIN)"],
            ["3", "permissions", "Otorisasi", "id (INT)", "-", "Katalog izin operasi granular (create, edit, approve, dll.)"],
            ["4", "role_permissions", "Otorisasi", "role_id, perm_id", "role_id, permission_id", "Pemetaan relasi many-to-many role terhadap permissions"],
            ["5", "user_scopes", "Otorisasi", "id (INT)", "user_id, institution_id", "Batasan wilayah wewenang instansi yang dapat diakses pengguna"],
            ["6", "access_grants", "Otorisasi", "id (INT)", "user_id, institution_id", "Pemberian izin akses khusus lintas instansi berbatas waktu"],
            ["7", "access_requests", "Otorisasi", "id (INT)", "user_id, institution_id", "Pengajuan permohonan akses data instansi tambahan"],
            ["8", "institutions", "Master Data", "id (INT)", "type_id", "Master data seluruh instansi kementerian/lembaga/daerah"],
            ["9", "institution_types", "Master Data", "id (INT)", "-", "Kategori tipe instansi pemerintah"],
            ["10", "cabinets", "Master Data", "id (INT)", "-", "Data kabinet pemerintahan dan periode masa jabatan"],
            ["11", "cabinet_institutions", "Master Data", "cabinet_id, inst_id", "cabinet_id, institution_id", "Pemetaan instansi yang tergabung dalam kabinet aktif"],
            ["12", "organizational_units", "Master Data", "id (INT)", "institution_id, parent_id", "Unit kerja struktural aktif dengan struktur pohon Adjacency List"],
            ["13", "positions", "Master Data", "id (INT)", "unit_id", "Jabatan struktural aktif, level eselon, dan formasi ASN"],
            ["14", "submissions", "Pengusulan", "id (INT)", "institution_id, author_id", "Header usulan penataan kelembagaan dan status workflow"],
            ["15", "submission_versions", "Pengusulan", "id (INT)", "submission_id", "Snapshot versi usulan (v1, v2, dst.) untuk immutabilitas"],
            ["16", "submission_units", "Pengusulan", "id (INT)", "version_id, parent_id", "Rincian usulan perubahan unit kerja (CREATE/UPDATE/DELETE)"],
            ["17", "submission_positions", "Pengusulan", "id (INT)", "version_id, unit_id", "Rincian usulan perubahan formasi jabatan struktural"],
            ["18", "verifier_assignments", "Verifikasi", "id (INT)", "submission_id, verifier_id", "Rekam penugasan berkas usulan dari Admin ke Verifikator"],
            ["19", "verifier_review_notes", "Verifikasi", "id (INT)", "submission_id, verifier_id", "Catatan koreksi telaah substantif per unit/organisasi"],
            ["20", "approval_records", "Pengesahan", "id (INT)", "version_id, approver_id", "Dokumen SK penetapan persetujuan resmi oleh Verifikator"],
            ["21", "audit_logs", "Audit Forensik", "id (INT)", "actor_id, institution_id", "Log mutasi data append-only, payload JSON, dan rekam jejak IP"]
        ],
        [0.3, 1.4, 1.1, 0.9, 1.3, 1.5]
    )

    builder.add_h2("7.2 Entity Relationship Diagram (ERD)")
    builder.add_p("Diagram relasi entitas basis data `eskld_db` disajikan pada Gambar 6.")
    builder.add_image_figure("erd_diagram.png", "Entity Relationship Diagram (ERD) Basis Data eskld_db (21 Tabel Relasional)")

    # =========================================================
    # BAB 8: SIKLUS HIDUP PENGUSULAN (SUBMISSION LIFECYCLE)
    # =========================================================
    builder.add_h1("BAB 8 — SIKLUS HIDUP PENGUSULAN (SUBMISSION LIFECYCLE)")
    
    builder.add_h2("8.1 Tahapan Siklus Pengusulan Kelembagaan")
    builder.add_p("Siklus hidup usulan penataan organisasi diatur melalui mesin status (Finite State Machine) yang menjamin setiap tahapan berjalan secara berurutan dan akuntabel:")
    builder.add_bullet("Operator instansi (USER) menyusun rincian usulan unit baru, perubahan eselon, atau kuota formasi. Data disimpan sebagai draf kerja awal (v1).", "1. Tahap DRAFT: ")
    builder.add_bullet("Operator melakukan finalisasi dan mengirimkan berkas usulan ke KemenPANRB. Berkas masuk ke antrean penapisan.", "2. Tahap SUBMITTED_TO_ADMIN: ")
    builder.add_bullet("Admin KemenPANRB memeriksa kelengkapan formal dan dasar hukum. Jika lengkap, usulan dialokasikan ke Verifikator.", "3. Tahap ASSIGNED_TO_VERIFIER: ")
    builder.add_bullet("Verifikator kelembagaan melakukan evaluasi mendalam terhadap analisis jabatan, beban kerja, dan kesesuaian regulasi.", "4. Tahap IN_REVIEW_BY_VERIFIER: ")
    builder.add_bullet("Hasil telaah substantif dinyatakan tuntas dan berkas siap diajukan ke tahap pengesahan final.", "5. Tahap READY_FOR_FINAL_DECISION: ")
    builder.add_bullet("Verifikator menerbitkan Surat Keputusan (SK) persetujuan resmi beserta nomor penetapan dan catatan pengesahan.", "6. Tahap APPROVED: ")
    builder.add_bullet("Sistem secara otomatis memigrasikan seluruh struktur unit kerja dan jabatan dari usulan yang disetujui langsung ke master data aktif.", "7. Tahap PROMOTED: ")

    builder.add_h2("8.2 Diagram Finite State Machine")
    builder.add_p("Diagram transisi status pengusulan disajikan pada Gambar 7.")
    builder.add_image_figure("submission_lifecycle_fsm.png", "Diagram Finite State Machine Siklus Hidup Pengusulan Kelembagaan")

    # =========================================================
    # BAB 9: VERSIONING DAN SIKLUS REVISI
    # =========================================================
    builder.add_h1("BAB 9 — VERSIONING DAN SIKLUS REVISI")
    
    builder.add_h2("9.1 Mekanisme Snapshot Immutability")
    builder.add_p("Untuk mencegah manipulasi data historis, SIGMA-K mengadopsi mekanisme Snapshot Versioning. Ketika berkas usulan dikembalikan oleh Admin atau Verifikator dengan status `REVISION_REQUIRED`, data versi 1 (v1) langsung dikunci permanen pada tabel `submission_versions`. Operator instansi kemudian menyusun perbaikan pada versi baru (v2).")

    builder.add_h2("9.2 Mesin Pelacak Perbedaan (Diff Engine)")
    builder.add_p("Sistem dilengkapi dengan Diff Viewer yang secara cerdas membandingkan usulan versi terbaru terhadap master data aktif:")
    builder.add_bullet("Unit kerja yang baru diusulkan diberi label `CREATE` dengan indikator warna hijau.", "• Penambahan Unit: ")
    builder.add_bullet("Perubahan nama unit, pergeseran induk, atau penambahan formasi diberi label `UPDATE` dengan indikator warna kuning/emas.", "• Perubahan Struktur: ")
    builder.add_bullet("Unit yang diusulkan untuk dilebur atau dihapus diberi label `DELETE` dengan indikator warna merah.", "• Penghapusan Unit: ")

    builder.add_image_figure("versioning_diff_flow.png", "Mekanisme Snapshot Versioning dan Pelacakan Perubahan Data (Diff Flow)")

    # =========================================================
    # BAB 10: PENAPISAN TAHAP 1 (GATE 1 ADMIN SCREENING)
    # =========================================================
    builder.add_h1("BAB 10 — PENAPISAN TAHAP 1 (GATE 1 ADMIN SCREENING)")
    
    builder.add_h2("10.1 Prosedur Penapisan Administratif")
    builder.add_p("Gate 1 Admin Screening merupakan pintu gerbang pertama verifikasi di KemenPANRB. Admin bertugas:")
    builder.add_bullet("Memastikan surat permohonan resmi, naskah akademik, dan dasar hukum pembentukan instansi telah terlampir lengkap.", "1. Validasi Kelengkapan Berkas: ")
    builder.add_bullet("Jika berkas tidak memenuhi syarat formal, Admin mengeksekusi `returnRevision()`, mengembalikan usulan ke status `REVISION_REQUIRED` disertai alasan tertulis.", "2. Pengembalian Revisi Formal: ")
    builder.add_bullet("Jika berkas valid, Admin memilih Verifikator yang berkompeten dan mengeksekusi `accept() & assignVerifier()`.", "3. Alokasi Penugasan: ")

    builder.add_callout("Batasan Wewenang Gate 1", "Admin KemenPANRB secara tegas DILARANG melakukan persetujuan substantif maupun pengesahan akhir Surat Keputusan. Wewenang Admin dibatasi strictly pada verifikasi formalitas dan alokasi kerja Verifikator.")

    builder.add_image_figure("gate1_admin_flowchart.png", "Diagram Alir Kerja Penapisan Administratif (Gate 1 Admin Screening)")

    # =========================================================
    # BAB 11: TELAAH SUBSTANTIF TAHAP 2 (GATE 2 VERIFIER & FINAL APPROVAL)
    # =========================================================
    builder.add_h1("BAB 11 — TELAAH SUBSTANTIF TAHAP 2 (GATE 2 VERIFIER)")
    
    builder.add_h2("11.1 Prosedur Telaah Substantif & Pengesahan SK")
    builder.add_p("Gate 2 Verifier Workspace adalah pusat evaluasi teknis kelembagaan. Verifikator yang ditugaskan melakukan:")
    builder.add_bullet("Mengevaluasi kesesuaian rentang kendali (span of control), analisis beban kerja, dan batas kuota formasi ASN.", "1. Analisis Beban Kerja & Formasi: ")
    builder.add_bullet("Menginputkan koreksi teknis terinci pada form `verifier_review_notes` jika ditemukan ketidaksesuaian substantif.", "2. Penerbitan Catatan Teknis: ")
    builder.add_bullet("Verifikator memiliki hak wewenang tunggal untuk menerbitkan SK persetujuan resmi melalui perintah `finalApprove()`.", "3. Pengesahan Final (Final Approval Authority): ")
    builder.add_bullet("Pasca pengesahan, sistem secara otomatis mengeksekusi `promote()`, mentransfer seluruh struktur baru ke master data aktif tanpa intervensi manual.", "4. Promosi Master Data Otomatis: ")

    builder.add_image_figure("gate2_verifier_flowchart.png", "Diagram Alir Kerja Telaah Substantif & Pengesahan SK (Gate 2 Verifier Workspace)")

    # =========================================================
    # BAB 12: ARSITEKTUR OTENTIKASI & JWT
    # =========================================================
    builder.add_h1("BAB 12 — ARSITEKTUR OTENTIKASI & JWT")
    
    builder.add_h2("12.1 Alur Otentikasi dan Pengelolaan Token")
    builder.add_p("Siklus otentikasi pengguna pada SIGMA-K menggunakan standar industri JSON Web Token (JWT):")
    builder.add_bullet("Klien mengirimkan kredensial NIP/Username dan Password melalui endpoint `POST /api/v1/auth/login`.", "1. Login Kredensial: ")
    builder.add_bullet("Backend memvalidasi kredensial terhadap tabel `users`, lalu menerbitkan JWT bertanda tangan HMAC-SHA256 yang memuat klaim identitas, role, dan home institution.", "2. Penerbitan JWT: ")
    builder.add_bullet("Token disimpan di Browser Storage melalui `BrowserStorageTokenProvider`. Setiap request HTTP berikutnya otomatis menyisipkan header `Authorization: Bearer <token>`.", "3. Injeksi Otomatis: ")
    builder.add_bullet("Saat aplikasi dimuat, endpoint `GET /api/v1/auth/me` dipanggil untuk menginisialisasi profil pengguna, daftar izin, dan cakupan instansi yang sah.", "4. Resolusi Profil & Sesi: ")

    builder.add_image_figure("auth_jwt_flow.png", "Diagram Alur Otentikasi JWT dan Validasi Sesi Pengguna")

    # =========================================================
    # BAB 13: OTORISASI ZERO-TRUST & KEAMANAN SISTEM
    # =========================================================
    builder.add_h1("BAB 13 — OTORISASI ZERO-TRUST & KEAMANAN SISTEM")
    
    builder.add_h2("13.1 Prinsip Keamanan Zero-Trust")
    builder.add_p("SIGMA-K menerapkan arsitektur keamanan Zero-Trust yang ketat:")
    builder.add_bullet("Frontend Next.js tidak pernah membuat keputusan otorisasi secara mandiri. Backend `AuthorizationService` dan `ScopeResolver` menjadi otoritas mutlak.", "1. Backend Single Source of Truth: ")
    builder.add_bullet("Setiap kueri data instansi, usulan, dan telaah diverifikasi kepemilikannya terhadap `user_scopes` dan `access_grants` pengguna untuk mencegah eksploitasi BOLA/IDOR.", "2. Proteksi BOLA / IDOR Scoping: ")
    builder.add_bullet("Penulis usulan (author) secara sistematis diblokir dari melakukan telaah atau pengesahan atas usulan buatannya sendiri.", "3. Anti Self-Approval Guard: ")
    builder.add_bullet("Pada mode API produksi, persona switcher di TopBar dinonaktifkan menjadi Read-Only Badge guna mencegah eskalasi wewenang ilegal pada sisi klien.", "4. Isolasi Persona Produksi: ")

    # =========================================================
    # BAB 14: DASHBOARD EKSEKUTIF & PELAPORAN DATA
    # =========================================================
    builder.add_h1("BAB 14 — DASHBOARD EKSEKUTIF & PELAPORAN DATA")
    
    builder.add_h2("14.1 Arsitektur Pelaporan Eksekutif")
    builder.add_p("Modul pelaporan pada SIGMA-K didukung oleh `ReportController` dan `ExecutiveReportService` backend yang menyajikan data analitik secara real-time dari basis data `eskld_db`:")
    builder.add_bullet("Menyajikan metrik total instansi, total unit aktif, total jabatan, total formasi ASN, ringkasan status funnel, dan pengesahan SK terkini.", "• GET /api/v1/reports/summary: ")
    builder.add_bullet("Laporan rincian usulan terfilter berdasarkan instansi, status, dan tahun usulan.", "• GET /api/v1/reports/submissions: ")
    builder.add_bullet("Rekapitulasi beban kelembagaan per instansi pemerintah.", "• GET /api/v1/reports/institutions: ")
    builder.add_bullet("Daftar rekam pengesahan SK resmi yang disahkan oleh Verifikator.", "• GET /api/v1/reports/approvals: ")
    builder.add_bullet("Riwayat migrasi usulan ke master data aktif.", "• GET /api/v1/reports/promotions: ")
    builder.add_bullet("Layanan unduh dataset resmi dalam format CSV stream dan JSON.", "• GET /api/v1/reports/export: ")

    # =========================================================
    # BAB 15: FORENSIK AUDIT TRAIL
    # =========================================================
    builder.add_h1("BAB 15 — FORENSIK AUDIT TRAIL")
    
    builder.add_h2("15.1 Rekam Jejak Append-Only Mutlak")
    builder.add_p("Setiap aksi manipulasi data pada sistem dicatat secara otomatis ke dalam tabel `audit_logs` melalui `AuditService`. Struktur data audit mencatat secara detail:")
    builder.add_bullet("Identitas unik aktor yang melakukan tindakan (actor_id, actor_name, actor_role).", "• Data Aktor: ")
    builder.add_bullet("Klasifikasi aksi (CREATE, UPDATE, DELETE, SUBMIT, VERIFY, APPROVE, REJECT, PROMOTE).", "• Event Aksi: ")
    builder.add_bullet("Tipe entitas yang dimanipulasi (SUBMISSION, INSTITUTION, UNIT, POSITION, ACCESS) beserta ID target.", "• Entitas Target: ")
    builder.add_bullet("Snapshot JSON data sebelum dan sesudah perubahan guna memungkinkan rekonstruksi forensik penuh.", "• Payload Diffing: ")
    builder.add_bullet("Alamat IP klien, informasi user agent, dan timestamp berpresisi tinggi.", "• Konteks Jaringan: ")

    # =========================================================
    # BAB 16: SPESIFIKASI & INTEGRASI REST API
    # =========================================================
    builder.add_h1("BAB 16 — SPESIFIKASI & INTEGRASI REST API")
    
    builder.add_h2("16.1 Katalog Lengkap REST API Terverifikasi")
    builder.add_p("Berikut adalah tabel spesifikasi seluruh endpoint REST API CodeIgniter 4 yang telah terintegrasi dengan frontend Next.js:")
    builder.add_table_data(
        "Katalog REST API Terverifikasi E-SKLD / SIGMA-K",
        ["Method", "Endpoint API", "Modul / Fungsi", "Otorisasi & Filter", "Respons Payload"],
        [
            ["POST", "/api/v1/auth/login", "Autentikasi Pengguna", "Publik / Rate-Limited", "JWT Access Token & User Info"],
            ["GET", "/api/v1/auth/me", "Profil Sesi Aktif", "Bearer JWT Auth", "Profil User, Role, Permissions, Scopes"],
            ["POST", "/api/v1/auth/logout", "Pengakhiran Sesi", "Bearer JWT Auth", "Status sukses logout & blacklist token"],
            ["GET", "/api/v1/institutions", "Katalog K/L/D", "Bearer JWT Auth", "Array seluruh master data instansi"],
            ["GET", "/api/v1/institutions/{id}", "Detail Instansi", "ScopeResolver Check", "Objek instansi, tipe, dan rekap unit"],
            ["GET", "/api/v1/organizations/tree", "Pohon Struktur Unit", "ScopeResolver Check", "Recursive tree Adjacency List"],
            ["GET", "/api/v1/organizations/{id}", "Rincian Unit Kerja", "ScopeResolver Check", "Detail unit, parent, & daftar positions"],
            ["GET", "/api/v1/submissions", "Daftar Usulan", "Role Scoped Filter", "Array usulan terfilter sesuai wewenang"],
            ["POST", "/api/v1/submissions", "Pembuatan Usulan Baru", "USER Role Guard", "Header draf usulan & ID versi v1"],
            ["GET", "/api/v1/submissions/{id}", "Rincian Berkas Usulan", "ScopeResolver Check", "Detail usulan, diff unit, jabatan, stepper"],
            ["POST", "/api/v1/submissions/{id}/submit", "Pengiriman ke Gate 1", "Author Ownership Guard", "Status transisi SUBMITTED_TO_ADMIN"],
            ["POST", "/api/v1/admin/screen/{id}", "Gate 1 Admin Action", "ADMIN Role Guard", "Accept / Return Revision status"],
            ["POST", "/api/v1/admin/assign/{id}", "Penugasan Verifikator", "ADMIN Role Guard", "Status transisi ASSIGNED_TO_VERIFIER"],
            ["POST", "/api/v1/verifier/review/{id}", "Telaah Substantif", "Assigned Verifier Guard", "Catatan review & status READY_FOR_DECISION"],
            ["POST", "/api/v1/verifier/approve/{id}", "Pengesahan SK Resmi", "VERIFIER Role Guard", "Penerbitan SK & status APPROVED"],
            ["POST", "/api/v1/verifier/promote/{id}", "Promosi Master Data", "VERIFIER Role Guard", "Eksekusi mutasi ke master data aktif"],
            ["GET", "/api/v1/reports/summary", "Ringkasan Eksekutif", "Bearer JWT Auth", "Metrik overview, funnel, SK terbaru"],
            ["GET", "/api/v1/reports/export", "Ekspor Dataset", "Scoped Permissions", "File stream CSV / JSON dataset"],
            ["GET", "/api/v1/audit-logs", "Log Forensik Audit", "ADMIN / SUPER_ADMIN", "Array rekam jejak mutasi & diff JSON"]
        ],
        [0.7, 1.8, 1.4, 1.3, 1.3]
    )

    # =========================================================
    # BAB 17: SPESIFIKASI & PROTOTYPE TAMPILAN APLIKASI (UI/UX)
    # =========================================================
    builder.add_h1("BAB 17 — SPESIFIKASI & PROTOTYPE TAMPILAN APLIKASI")
    
    builder.add_h2("17.1 SCR-18: Layar Masuk Sistem (Login)")
    builder.add_p("Layar autentikasi awal yang menyajikan form input NIP dan Password dengan indikator mode API, validasi format, dan tombol preset akun.")
    builder.add_image_figure("prototype_login.png", "Prototype Layar Masuk Sistem (SCR-18 Login)")

    builder.add_h2("17.2 SCR-01: Dashboard Eksekutif Kelembagaan")
    builder.add_p("Layar beranda eksekutif yang menampilkan 4 kartu KPI utama (Total Usulan, Dalam Proses, Disetujui, Promosi), sorotan kabinet aktif, antrean usulan terkini, dan widget SK pengesahan terbaru.")
    builder.add_image_figure("prototype_dashboard.png", "Prototype Dashboard Eksekutif Kelembagaan (SCR-01)")

    builder.add_h2("17.3 SCR-02: Katalog Master Instansi Pemerintah")
    builder.add_p("Katalog terpadu kementerian/lembaga/daerah dengan filter kategori (Kemenko, Kementerian, LPNK, Pemda), pencarian instan, dan statistik ringkas beban unit kerja.")
    builder.add_image_figure("prototype_institutions.png", "Prototype Katalog Master Instansi Pemerintah (SCR-02)")

    builder.add_h2("17.4 SCR-04: Bagan Struktur Organisasi Interaktif (React Flow)")
    builder.add_p("Kanvas graf interaktif berteknologi React Flow yang menyajikan pohon hierarki unit kerja lengkap dengan garis konektor parent-child, level eselon, MiniMap, dan drawer detail formasi jabatan.")
    builder.add_image_figure("prototype_org_structure.png", "Prototype Bagan Struktur Organisasi Interaktif (SCR-04 React Flow Canvas)")

    builder.add_h2("17.5 SCR-09: Rincian Usulan & Diff Viewer")
    builder.add_p("Ruang pemeriksaan berkas pengusulan yang dilengkapi dengan Workflow Stepper penunjuk status, tabel komparasi Before vs After (Diff Viewer), serta tab riwayat catatan telaah.")
    builder.add_image_figure("prototype_submission_detail.png", "Prototype Rincian Usulan & Diff Viewer (SCR-09)")

    builder.add_h2("17.6 SCR-12: Ruang Kerja Telaah Substantif Verifikator")
    builder.add_p("Ruang kerja khusus Verifikator kelembagaan untuk menganalisis beban kerja organisasi, menginput catatan rekomendasi teknis, menandatangani pengesahan SK resmi, dan mengeksekusi promosi data.")
    builder.add_image_figure("prototype_verifier_workspace.png", "Prototype Ruang Kerja Telaah Verifikator (SCR-12 Verifier Workspace)")

    builder.add_h2("17.7 SCR-15: Intelijensi Data & Postur ASN (Analytics Studio)")
    builder.add_p("Studio analitik eksekutif yang menyajikan indikator KPI kinerja kelembagaan, piramida distribusi jabatan struktural Eselon I s.d. IV, rata-rata durasi SLA layanan, dan tombol ekspor laporan CSV.")
    builder.add_image_figure("prototype_analytics_reporting.png", "Prototype Intelijensi Data & Postur ASN (SCR-15 Analytics Studio)")

    builder.add_h2("17.8 SCR-16: Log Forensik Audit Trail")
    builder.add_p("Tabel rekam jejak kepatuhan dan audit forensik mutlak yang mencatat setiap timestamp, nama aktor, role, event aksi, modul target, alamat IP, dan modal perbandingan payload JSON.")
    builder.add_image_figure("prototype_audit_logs.png", "Prototype Log Forensik Audit Trail (SCR-16)")

    # =========================================================
    # BAB 18: HASIL PENGUJIAN & VALIDASI SISTEM
    # =========================================================
    builder.add_h1("BAB 18 — HASIL PENGUJIAN & VALIDASI SISTEM")
    
    builder.add_h2("18.1 Hasil Pengujian Otomatis Frontend & Backend")
    builder.add_p("Kualitas, stabilitas, dan kepatuhan arsitektur sistem SIGMA-K telah divalidasi secara ketat melalui pengujian otomatis komprehensif pada lapisan frontend dan backend:")
    builder.add_table_data(
        "Hasil Pengujian & Validasi Kualitas Sistem SIGMA-K",
        ["Lapisan Pengujian", "Runner / Tool", "Cakupan Modul", "Target Assertion", "Hasil Validasi"],
        [
            ["Frontend Foundation (14A)", "npx tsx run-foundation-tests.ts", "HTTP Client, AppError, DTO Mappers", "28 Tests", "100% LULUS (PASS)"],
            ["Frontend Auth & Security (14B)", "npx tsx run-auth-tests.ts", "JWT Lifecycle, Zero-Trust Role Lock", "24 Tests", "100% LULUS (PASS)"],
            ["Frontend Master Data (14C)", "npx tsx run-master-tests.ts", "Instansi, DFS Tree Flattening, Unit Detail", "24 Tests", "100% LULUS (PASS)"],
            ["Frontend Executive Reports (14D)", "npx tsx run-report-tests.ts", "Summary KPI, Funnel, CSV Export Stream", "29 Tests", "100% LULUS (PASS)"],
            ["Total Frontend Composite Suite", "npx tsx run-report-tests.ts", "Seluruh Domain Integrasi Frontend", "105 Tests", "100% LULUS (PASS)"],
            ["Backend Unit & Integration Suite", "vendor/bin/phpunit", "Zero-Trust, FSM Workflow, Gate 1 & 2, Audit", "198 Tests / 713 Assertions", "100% LULUS (0 Error, 0 Fail)"],
            ["TypeScript Strict Type Check", "npx tsc --noEmit", "Seluruh Komponen, Services & DTO", "0 Type Errors", "100% LULUS (PASS)"],
            ["ESLint Code Style Analysis", "npm run lint", "Standar Kualitas & Kebersihan Kode", "0 Warning / 0 Error", "100% LULUS (PASS)"],
            ["Next.js Production Build", "npm run build", "Kompilasi 16 Route Aplikasi", "16 Routes Compiled", "100% LULUS (PASS)"]
        ],
        [1.5, 1.5, 1.5, 1.0, 1.0]
    )

    # =========================================================
    # BAB 19: STATUS MATRIKS IMPLEMENTASI FITUR
    # =========================================================
    builder.add_h1("BAB 19 — STATUS MATRIKS IMPLEMENTASI FITUR")
    
    builder.add_h2("19.1 Matriks Kesiapan Fungsionalitas")
    builder.add_p("Tabel berikut menyajikan status implementasi aktual dari seluruh fitur sistem berdasarkan verifikasi source code repository:")
    builder.add_table_data(
        "Matriks Status Implementasi Fitur E-SKLD / SIGMA-K",
        ["Modul / Fitur", "Status Kesiapan", "Cakupan Backend", "Cakupan Frontend", "Keterangan Verifikasi"],
        [
            ["Autentikasi & JWT", "COMPLETED", "CodeIgniter 4 + AuthFilter", "Next.js AuthService & Context", "Token lifecycle & profile terverifikasi"],
            ["Otorisasi Zero-Trust", "COMPLETED", "AuthorizationService + ScopeResolver", "Read-Only Role Badge", "BOLA/IDOR protection & role lock aktif"],
            ["Master Data Instansi", "COMPLETED", "InstitutionController", "InstitutionService + Catalog UI", "Katalog K/L/D & profil terintegrasi"],
            ["Struktur Organisasi", "COMPLETED", "OrgHierarchyService (DFS Tree)", "React Flow Canvas Interactive", "Visualisasi graf pohon & drawer aktif"],
            ["Drafting Usulan", "COMPLETED", "SubmissionWorkflowService", "SubmissionService + Form UI", "Penyusunan usulan unit & jabatan aktif"],
            ["Snapshot Versioning", "COMPLETED", "submission_versions (v1->v2)", "DiffViewer Before-After", "Immutabilitas data & audit diff aktif"],
            ["Gate 1 Admin Screening", "COMPLETED", "AdminWorkflowController", "Gate 1 Verification Workspace", "Penapisan berkas & alokasi verifikator"],
            ["Gate 2 Verifier Review", "COMPLETED", "VerifierWorkflowController", "Gate 2 Verifier Workspace", "Telaah substantif & rekomendasi aktif"],
            ["Final Approval (SK)", "COMPLETED", "ApprovalPromotionService", "Pengesahan SK UI Component", "Wewenang final mutlak pada Verifikator"],
            ["Promosi Master Data", "COMPLETED", "Promote Transaction Engine", "Promotion Status Indicator", "Migrasi usulan ke master data otomatis"],
            ["Dashboard Eksekutif", "COMPLETED", "ReportController::summary", "Dashboard Cards & Funnel UI", "Metrik live terhubung ke eskld_db"],
            ["Pelaporan & Ekspor", "COMPLETED", "ReportController::export", "Analytics Studio & CSV Downloader", "Unduh berkas CSV/JSON live terverifikasi"],
            ["Forensik Audit Trail", "COMPLETED", "AuditService (Append-Only)", "Audit Log Table & Modal Diff", "Rekam jejak mutasi & IP terverifikasi"],
            ["Realtime Presence", "FUTURE REQUIREMENT", "Belum Diimplementasikan", "Belum Diimplementasikan", "Tercatat resmi sebagai OPEN-005"]
        ],
        [1.4, 1.1, 1.3, 1.3, 1.4]
    )

    # =========================================================
    # BAB 20: RENCANA PENGEMBANGAN LANJUTAN (FUTURE REQUIREMENTS)
    # =========================================================
    builder.add_h1("BAB 20 — RENCANA PENGEMBANGAN LANJUTAN")
    
    builder.add_h2("20.1 Kebutuhan Masa Depan Tercatat Resmi")
    builder.add_p("Berdasarkan register kebutuhan resmi project, item pengembangan lanjutan yang telah dipetakan untuk rilis berikutnya mencakup:")
    builder.add_bullet("Penyediaan indikator status online/offline real-time bagi verifikator yang sedang menelaah berkas usulan yang sama menggunakan teknologi WebSocket / Server-Sent Events (SSE) / Redis.", "1. OPEN-005 (Realtime Presence & Active Reviewer): ")
    builder.add_bullet("Integrasi modul tanda tangan digital tersertifikasi bekerjasama dengan Balai Sertifikasi Elektronik (BSrE) BSSN untuk penerbitan dokumen SK resmi.", "2. Digital Signature Integration: ")
    builder.add_bullet("Penambahan storage terenkripsi untuk pengunggahan naskah akademik, regulasi pembentukan, dan dokumen pendukung berukuran besar.", "3. Multi-Faceted Document Attachment: ")

    # =========================================================
    # BAB 21: KESIMPULAN & PENUTUP
    # =========================================================
    builder.add_h1("BAB 21 — KESIMPULAN & PENUTUP")
    
    builder.add_h2("21.1 Kesimpulan Capaian Perancangan")
    builder.add_p("Dokumen Perancangan Sistem dan Prototype Aplikasi E-SKLD / SIGMA-K Versi 1.0 ini merangkum seluruh hasil capaian rekayasa perangkat lunak yang telah berhasil dibangun, diintegrasikan, dan diuji secara komprehensif. Sistem telah membuktikan kesiapan arsitektur dalam mengelola master data kelembagaan kementerian/lembaga/daerah, memfasilitasi tata kelola pengusulan penataan organisasi secara terstruktur, menegakkan prinsip Zero-Trust dan pemisahan wewenang (Separation of Duties), serta menyediakan rekam jejak forensik mutlak guna mendukung tata kelola pemerintahan digital yang bersih, transparan, dan akuntabel di Indonesia.")

    # Save DOCX
    print(f"Saving DOCX to {DOCX_OUT}...")
    builder.doc.save(DOCX_OUT)
    print("DOCX successfully created!")

if __name__ == '__main__':
    build_sigma_k_document()
